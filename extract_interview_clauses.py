import os
import docx
import spacy
import openpyxl
from pathlib import Path
import re


# --- Editable lexicons (tweak these as you iterate) ---
OPINION_VERBS = {
    "like",
    "love",
    "enjoy",
    "prefer",
    "appreciate",
    "hate",
    "dislike",
    "want",
    "need",
    "desire",
}

# Evaluative verbs: verbs that typically encode a benefit/cost judgement where the
# evaluated "object" is often the subject/action itself (e.g., "X saves time").
EVALUATIVE_VERBS = {
    "save",
    "help",
    "improve",
}

ENABLE_VERBS = {
    "allow",
    "enable",
    "permit",
    "let",
    "get",
    "try",
    "can",
    "pick",
}

COMPARATIVE_WORDS = {
    "fewer",
    "less",
    "more",
    "better",
    "worse",
    "best",
    "worst",
    "easier",
    "quicker",
    "harder",
    "slower",
    "faster",
}

EVALUATIVE_ADJECTIVES = {
    "quick",
    "consistent",
    "slow",
    "good",
    "bad",
    "great",
    "poor",
    "helpful",
    "unhelpful",
    "clear",
    "unclear",
    "detailed",
    "vague",
    "easy",
    "difficult",
    "readable",
    "unreadable",
    "legible",
    "illegible",
    "absolute",
    "much",
    "correct",
    "incorrect",
}

EVALUATIVE_NOUNS = {
    "difficulty",
    "trouble",
    "problem",
    "issue",
    "challenge",
    "benefit",
    "advantage",
    "disadvantage",
    "improvement",
    "irritation",
    "fan",
    "fantastic",
    "clear",
    "favour",
    "nice",
    "ability",
    "lot",
    "concern",
    "issue",
    "bias",
    "ease",
}

EVALUATIVE_ADVERBS = {
    "definitely",
    "definately",  # common misspelling
    "certainly",
    "really",
    "quite",
    "very",
    "clearly",
    "obviously",
    "completely",
    "really",
    "automatically",
    "strongly",
    "correctly",
    "incorrectly",
    "easily",
    "properly",
    "so much",
}

# Load spaCy model
def load_spacy_model():
    """Load a spaCy English pipeline.

    Uses en_core_web_lg (better parses). Downloads it if missing.
    """

    try:
        return spacy.load("en_core_web_lg")
    except OSError:
        from spacy.cli import download

        download("en_core_web_lg")
        return spacy.load("en_core_web_lg")


def _find_of_head_noun(token):
    """If token occurs inside an 'of' phrase attached to a noun, return that noun.

    Example: "examples of them having difficulty" -> for 'difficulty', returns 'examples'.
    """

    for anc in (token, *token.ancestors):
        # token inside pobj/pcomp of an 'of' preposition
        if anc.dep_ in {"pobj", "pcomp"} and anc.head.dep_ == "prep" and anc.head.lemma_.lower() == "of":
            head = anc.head.head
            if head is not None and head.pos_ in {"NOUN", "PROPN"}:
                return head
        # token is itself under an 'of' preposition
        if anc.dep_ == "prep" and anc.lemma_.lower() == "of":
            head = anc.head
            if head is not None and head.pos_ in {"NOUN", "PROPN"}:
                return head
    return None

def classify_adv(tok):
    lemma = tok.lemma_.lower()

    if lemma in {"always","usually","often","sometimes","rarely","never","frequently"}:
        return "FREQUENCY_ADV"
    if lemma in {"obviously","clearly","evidently","apparently"}:
        return "EVIDENCE_ADV"

    if lemma == "then":
        # look left in same sentence for an "if" marker
        sent = tok.sent
        left = tok.doc[sent.start:tok.i]
        has_if_mark = any(t.lower_ == "if" and t.dep_ == "mark" for t in left)

        prev_is_punct = tok.i > sent.start and tok.nbor(-1).is_punct

        if has_if_mark:
            return "COND_THEN"   # “if ..., then ...”
        if prev_is_punct:
            return "DISCOURSE_THEN"
        return "TEMPORAL_THEN"

    return "OTHER_ADV"


def _noun_phrase_span(token):
    if token is None:
        return None
    left = token.left_edge.i
    right = token.right_edge.i
    return token.doc[left : right + 1]


def _object_text_from_noun_head(head, *, evaluative_adjectives, comparative_words, exclude_tokens=None):
    if head is None:
        return ""
    exclude_tokens = set(exclude_tokens or [])

    tokens = []
    for tok in head.subtree:
        if tok in exclude_tokens:
            continue
        if tok.is_punct or tok.pos_ == "PUNCT":
            continue
        if tok.pos_ in {"INTJ"}:
            continue
        if tok.dep_ in {"det"}:
            continue
        if tok.dep_ in {"cc", "punct", "mark"}:
            continue
        if tok.dep_ == "conj" and tok is not head:
            # Avoid pulling in coordinated NPs into a single object string.
            continue
        if tok.pos_ == "ADV" or tok.dep_ == "advmod":
            continue

        # Drop comparative quantifiers from the object phrase (e.g., fewer/less/more).
        # Note: spaCy lemmas for comparatives can be base forms (e.g., "fewer" -> "few").
        if tok.lower_ in comparative_words or tok.lemma_.lower() in comparative_words:
            continue

        # Drop evaluative/comparative adjectival modifiers (keep non-evaluative ones like "admin" if tagged as ADJ)
        if tok.dep_ == "amod":
            lemma = tok.lemma_.lower()
            if lemma in evaluative_adjectives or lemma in comparative_words or tok.lower_ in comparative_words:
                continue

        tokens.append(tok)

    tokens = sorted(tokens, key=lambda t: t.i)
    return " ".join(t.text for t in tokens).strip()


def _verb_object_span(verb_token):
    if verb_token is None:
        return None

    # Prefer direct object
    obj = next((c for c in verb_token.children if c.dep_ in {"dobj", "obj"}), None)
    if obj is not None:
        return _noun_phrase_span(obj)

    # Copular/predicate constructions: "X is Y" where Y may be attr/acomp
    pred = next((c for c in verb_token.children if c.dep_ in {"attr", "acomp", "oprd"}), None)
    if pred is not None:
        return _noun_phrase_span(pred)

    # Prepositional object: look for prep -> pobj
    prep = next((c for c in verb_token.children if c.dep_ == "prep"), None)
    if prep is not None:
        pobj = next((c for c in prep.children if c.dep_ == "pobj"), None)
        if pobj is not None:
            return _noun_phrase_span(pobj)

    # Clausal complement
    comp = next((c for c in verb_token.children if c.dep_ in {"ccomp", "xcomp"}), None)
    if comp is not None:
        return comp.doc[comp.left_edge.i : comp.right_edge.i + 1]

    return None


def extract_evaluations(text, nlp):
    """Return a list of (evaluative_term, object_text) pairs for a clause.

    Heuristic and editable: uses small lexicons + dependency links.
    """

    doc = nlp(text)

    opinion_verbs = OPINION_VERBS
    evaluative_verbs = EVALUATIVE_VERBS
    comparative_words = COMPARATIVE_WORDS
    evaluative_adjectives = EVALUATIVE_ADJECTIVES
    enable_verbs = ENABLE_VERBS
    evaluative_nouns = EVALUATIVE_NOUNS
    evaluative_adverbs = EVALUATIVE_ADVERBS

    pairs = []

    def object_for_modifier(head_token):
        """Infer evaluated object for a modifier (ADV/ADJ) by walking its head chain."""

        if head_token is None:
            return ""

        tok = head_token
        for _ in range(4):
            # ADJ attached to a NOUN (covers amod and comparative determiners like "fewer")
            if tok.pos_ == "ADJ" and tok.head is not None and tok.head.pos_ in {"NOUN", "PROPN"}:
                return _object_text_from_noun_head(
                    tok.head,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )

            # ADJ modifying NOUN: quick grading / fewer queries
            if tok.pos_ == "ADJ" and tok.dep_ == "amod" and tok.head is not None and tok.head.pos_ in {"NOUN", "PROPN"}:
                return _object_text_from_noun_head(
                    tok.head,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )

            # Predicate ADJ: scans are unreadable
            if tok.pos_ == "ADJ" and tok.dep_ in {"acomp", "attr"}:
                cop = tok.head
                subj = next((c for c in cop.children if c.dep_ in {"nsubj", "nsubjpass"}), None)
                if subj is not None and subj.pos_ in {"NOUN", "PROPN", "PRON"}:
                    return _object_text_from_noun_head(
                        subj,
                        evaluative_adjectives=evaluative_adjectives,
                        comparative_words=comparative_words,
                    )

            # VERB: use its object span
            if tok.pos_ == "VERB":
                span = _verb_object_span(tok)
                if span is None:
                    return ""

                noun_heads = [t for t in span if t.pos_ in {"NOUN", "PROPN"}]
                if noun_heads:
                    head = noun_heads[-1]
                    return _object_text_from_noun_head(
                        head,
                        evaluative_adjectives=evaluative_adjectives,
                        comparative_words=comparative_words,
                    )

                return span.text.strip()

            # NOUN: treat noun itself as object
            if tok.pos_ in {"NOUN", "PROPN"}:
                return _object_text_from_noun_head(
                    tok,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )

            if tok.head is None or tok.head is tok:
                break
            tok = tok.head

        return ""

    def add_pair(term, obj_text):
        if not term:
            return
        obj_text = (obj_text or "").strip()
        term_text = term.strip()
        key = (term_text.lower(), obj_text.lower())
        if not hasattr(add_pair, "seen"):
            add_pair.seen = set()  # type: ignore[attr-defined]
        if key in add_pair.seen:  # type: ignore[attr-defined]
            return
        add_pair.seen.add(key)  # type: ignore[attr-defined]
        pairs.append((term_text, obj_text))

    # 1) Opinion/evaluative verbs
    for tok in doc:
        if tok.pos_ != "VERB":
            continue
        if tok.lemma_.lower() not in opinion_verbs:
            continue

        neg = any(c.dep_ == "neg" for c in tok.children)
        term = f"not {tok.lemma_}" if neg else tok.lemma_
        obj_head = next((c for c in tok.children if c.dep_ in {"dobj", "obj"} and c.pos_ in {"NOUN", "PROPN"}), None)
        if obj_head is not None:
            obj_heads = [obj_head, *list(obj_head.conjuncts)]
            obj_texts = []
            for h in obj_heads:
                obj_texts.append(
                    _object_text_from_noun_head(
                        h,
                        evaluative_adjectives=evaluative_adjectives,
                        comparative_words=comparative_words,
                    )
                )
            obj_texts = [t for t in obj_texts if t]
            obj_text = "; ".join(dict.fromkeys(obj_texts))
        else:
            # Fallback to wider patterns
            obj_span = _verb_object_span(tok)
            obj_text = obj_span.text.strip() if obj_span is not None else ""
        add_pair(term, obj_text)

    # 1a) Evaluative verbs like "save" where the evaluated object is typically the subject
    # (e.g., "Not having to ... saves time" -> term: "saves", object: "Not having to ...").
    for tok in doc:
        if tok.pos_ != "VERB":
            continue
        if tok.lemma_.lower() not in evaluative_verbs:
            continue

        neg = any(c.dep_ == "neg" for c in tok.children)
        term = tok.text if not neg else f"not {tok.lemma_}"

        # If the evaluative verb is transitive, prefer its object ("improve your speed").
        # Otherwise, fall back to the subject span for gerund/infinitival subjects
        # ("Not having to ... saves time").
        obj = next((c for c in tok.children if c.dep_ in {"dobj", "obj"}), None)
        if obj is not None:
            if obj.pos_ in {"NOUN", "PROPN"}:
                obj_text = _object_text_from_noun_head(
                    obj,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )
            else:
                span = _noun_phrase_span(obj)
                obj_text = span.text.strip() if span is not None else obj.text

            # Remove leading possessive pronouns for cleaner object labels ("your speed" -> "speed").
            obj_text = re.sub(r"^(my|your|his|her|our|their|its)\s+", "", obj_text, flags=re.IGNORECASE)
        else:
            subj = next((c for c in tok.children if c.dep_ in {"nsubj", "nsubjpass", "csubj"}), None)
            if subj is not None:
                span = _noun_phrase_span(subj)
                obj_text = span.text.strip() if span is not None else subj.text
            else:
                obj_text = ""

        add_pair(term, obj_text)

    # 1b) Enable/ability verbs ("allow/enable/permit/let")
    for tok in doc:
        if tok.pos_ != "VERB":
            continue
        if tok.lemma_.lower() not in enable_verbs:
            continue

        neg = any(c.dep_ == "neg" for c in tok.children)
        term = f"not {tok.lemma_}" if neg else tok.lemma_

        # Prefer infinitival/clausal complement as the "enabled action"
        comp = next((c for c in tok.children if c.dep_ in {"xcomp", "ccomp"}), None)
        if comp is not None:
            # Special case: "get to VERB ..." -> evaluative term "get to VERB" and object from VERB's object.
            if tok.lemma_.lower() == "get" and comp.pos_ == "VERB":
                term = f"{'not ' if neg else ''}get to {comp.lemma_}"

                obj_head = next((c for c in comp.children if c.dep_ in {"dobj", "obj"} and c.pos_ in {"NOUN", "PROPN"}), None)
                if obj_head is not None:
                    obj_text = _object_text_from_noun_head(
                        obj_head,
                        evaluative_adjectives=evaluative_adjectives,
                        comparative_words=comparative_words,
                    )
                else:
                    # Fall back to a cleaned verb object span; strip leading "to" if present.
                    span = _verb_object_span(comp)
                    obj_text = span.text.strip() if span is not None else ""
                    obj_text = re.sub(r"^to\s+", "", obj_text, flags=re.IGNORECASE)
            else:
                obj_text = comp.doc[comp.left_edge.i : comp.right_edge.i + 1].text.strip()
        else:
            # Or a direct object NP
            obj = next((c for c in tok.children if c.dep_ in {"dobj", "obj"}), None)
            obj_text = _noun_phrase_span(obj).text.strip() if obj is not None else ""

        add_pair(term, obj_text)

    # 1c) Modal ability/enable AUX: "can ..." (e.g., "comments can be used")
    for tok in doc:
        if tok.pos_ != "AUX":
            continue
        if tok.lemma_.lower() != "can":
            continue

        # In spaCy, modal AUX typically attaches to the main verb as an aux.
        verb = tok.head
        if verb is None:
            continue
        if verb.pos_ != "VERB":
            verb = next((a for a in tok.ancestors if a.pos_ == "VERB"), None)
        if verb is None or verb.pos_ != "VERB":
            continue

        # Build evaluative term text like "can be used".
        term_tokens = [tok]
        for child in verb.children:
            if child.dep_ in {"aux", "auxpass"} and child.pos_ == "AUX" and tok.i <= child.i <= verb.i:
                term_tokens.append(child)
        for child in (*tok.children, *verb.children):
            if child.dep_ == "neg" and tok.i <= child.i <= verb.i:
                term_tokens.append(child)
        term_tokens.append(verb)

        term_tokens = sorted({t.i: t for t in term_tokens}.values(), key=lambda t: t.i)
        term_text = " ".join(t.text for t in term_tokens)
        term_text = re.sub(r"\bca\s+n't\b", "can't", term_text, flags=re.IGNORECASE)
        term_text = re.sub(r"\s+", " ", term_text).strip()

        subj = next((c for c in verb.children if c.dep_ in {"nsubj", "nsubjpass"}), None)
        if subj is not None:
            span = _noun_phrase_span(subj)
            obj_text = span.text.strip() if span is not None else subj.text
        else:
            obj_span = _verb_object_span(verb)
            obj_text = obj_span.text.strip() if obj_span is not None else ""

        add_pair(term_text, obj_text)

    # 2) Evaluative/comparative adjectives modifying a noun (amod)
    for tok in doc:
        if tok.dep_ != "amod":
            continue
        if tok.pos_ != "ADJ":
            continue

        lemma = tok.lemma_.lower()
        is_comparative = tok.tag_ in {"JJR", "JJS"} or lemma in comparative_words
        is_eval_adj = lemma in evaluative_adjectives

        # Heuristic: treat "too + ADJ" as evaluative even if ADJ isn't in the lexicon.
        has_too = any(c.dep_ == "advmod" and c.lemma_.lower() == "too" for c in tok.children)

        if not (is_comparative or is_eval_adj or has_too):
            continue

        head = tok.head
        if head.pos_ not in {"NOUN", "PROPN"}:
            continue
        term_text = tok.text
        if has_too:
            term_text = f"too {term_text}"

        obj_text = _object_text_from_noun_head(
            head,
            evaluative_adjectives=evaluative_adjectives,
            comparative_words=comparative_words,
            exclude_tokens={tok},
        )
        add_pair(term_text, obj_text)

    # 2b) Predicate adjectives ("X is unreadable" / "that was helpful")
    for tok in doc:
        if tok.pos_ != "ADJ":
            continue
        # Include ADJ complements under verbs like "make" where spaCy can attach
        # the adjective as a clausal complement (e.g., "makes it more consistent").
        if tok.dep_ not in {"acomp", "attr", "oprd", "xcomp", "ccomp"}:
            continue

        lemma = tok.lemma_.lower()
        degree_mod = next(
            (c for c in tok.children if c.dep_ == "advmod" and c.lemma_.lower() in {"more", "less"}),
            None,
        )

        is_comparative = tok.tag_ in {"JJR", "JJS"} or lemma in comparative_words or degree_mod is not None
        is_eval_adj = lemma in evaluative_adjectives
        has_too = any(c.dep_ == "advmod" and c.lemma_.lower() == "too" for c in tok.children)
        if not (is_comparative or is_eval_adj or has_too):
            continue

        # Object inference:
        # - For copular: take subject of the copular verb.
        # - For "make/keep/render"-style complements: often the object is the ADJ's subject
        #   (spaCy often makes the object PRON a nsubj of the ADJ).
        head = tok.head
        subj_text = ""

        adj_subj = next((c for c in tok.children if c.dep_ in {"nsubj", "nsubjpass", "csubj"}), None)
        if adj_subj is not None:
            span = _noun_phrase_span(adj_subj)
            subj_text = span.text.strip() if span is not None else adj_subj.text
        elif head is not None and head.lemma_.lower() == "be":
            cop_subj = next((c for c in head.children if c.dep_ in {"nsubj", "nsubjpass"}), None)
            if cop_subj is not None:
                subj_text = _object_text_from_noun_head(
                    cop_subj,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )
        elif head is not None and head.pos_ == "VERB":
            obj = next((c for c in head.children if c.dep_ in {"dobj", "obj"}), None)
            if obj is not None:
                span = _noun_phrase_span(obj)
                subj_text = span.text.strip() if span is not None else obj.text

        term_text = tok.text
        if degree_mod is not None and degree_mod.i < tok.i:
            term_text = f"{degree_mod.text} {term_text}"
        if has_too:
            term_text = f"too {tok.text}"
        add_pair(term_text, subj_text)

    # 2c) Ability construction: "be able to ..." (treat as evaluative/ability term)
    for tok in doc:
        if tok.pos_ != "ADJ" or tok.lemma_.lower() != "able":
            continue
        if tok.dep_ not in {"acomp", "attr"}:
            continue

        head = tok.head

        # Negation can attach to the copular verb or sometimes to the ADJ
        neg = any(c.dep_ == "neg" for c in head.children) or any(c.dep_ == "neg" for c in tok.children)
        term_text = "be able" if not neg else "not be able"

        # Find the enabled action (xcomp/ccomp). Depending on the parse, this can hang off
        # the copular verb ("was") or off the adjective ("able").
        comp = next((c for c in tok.children if c.dep_ in {"xcomp", "ccomp"}), None)
        if comp is None:
            comp = next((c for c in head.children if c.dep_ in {"xcomp", "ccomp"}), None)
        if comp is not None:
            obj_text = comp.doc[comp.left_edge.i : comp.right_edge.i + 1].text.strip()
        else:
            # Fallback: the subject ("I was able")
            subj = next((c for c in head.children if c.dep_ in {"nsubj", "nsubjpass"}), None)
            if subj is not None and subj.pos_ in {"NOUN", "PROPN", "PRON"}:
                obj_text = _noun_phrase_span(subj).text.strip()
            else:
                obj_text = ""

        add_pair(term_text, obj_text)

    # 3) Comparative determiners like "fewer" directly attached to a noun
    for tok in doc:
        lemma = tok.lemma_.lower()
        if lemma not in comparative_words:
            continue
        if tok.dep_ not in {"det", "amod"}:
            continue
        head = tok.head
        if head.pos_ not in {"NOUN", "PROPN"}:
            continue
        obj_text = _object_text_from_noun_head(
            head,
            evaluative_adjectives=evaluative_adjectives,
            comparative_words=comparative_words,
            exclude_tokens={tok},
        )
        add_pair(tok.text, obj_text)

    # 4) Evaluative nouns ("difficulty", "problem", "issue" ...)
    for tok in doc:
        if tok.pos_ != "NOUN":
            continue
        if tok.lemma_.lower() not in evaluative_nouns:
            continue

        term_text = tok.text
        det_no = next((c for c in tok.children if c.dep_ == "det" and c.lemma_.lower() == "no"), None)
        if det_no is not None:
            term_text = f"no {tok.lemma_}"

        # If noun takes a PP complement ("difficulty with X"), treat X as the object.
        prep = next((c for c in tok.children if c.dep_ == "prep"), None)
        if prep is not None:
            pobj = next((c for c in prep.children if c.dep_ == "pobj"), None)
            if pobj is not None and pobj.pos_ in {"NOUN", "PROPN", "PRON"}:
                obj_text = _object_text_from_noun_head(
                    pobj,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )
                add_pair(term_text, obj_text)
                continue

        # Prefer the head noun of an enclosing "of" phrase: "examples of ... having difficulty".
        of_head = _find_of_head_noun(tok)
        if of_head is not None:
            obj_text = _object_text_from_noun_head(
                of_head,
                evaluative_adjectives=evaluative_adjectives,
                comparative_words=comparative_words,
            )
            add_pair(term_text, obj_text)
            continue

        # If the evaluative noun is an object of a verb ("have difficulty"), use the verb's subject.
        if tok.dep_ in {"dobj", "obj"} and tok.head.pos_ == "VERB":
            subj = next((c for c in tok.head.children if c.dep_ in {"nsubj", "nsubjpass"}), None)
            if subj is not None and subj.pos_ in {"NOUN", "PROPN", "PRON"}:
                obj_text = _object_text_from_noun_head(
                    subj,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                )
                add_pair(term_text, obj_text)
                continue

        # Fallback: no clear object
        add_pair(term_text, "")

    # 5) Evaluative/stance adverbs ("definitely", "certainly", etc.)
    for tok in doc:
        if tok.pos_ != "ADV" or tok.dep_ != "advmod":
            continue

        # Multi-word stance: "(not) so much" as in "Gradescope, not so much."
        # spaCy often parses: so(ADV advmod -> much) and not(PART neg -> much).
        if tok.lemma_.lower() == "so" and tok.head is not None and tok.head.lemma_.lower() == "much":
            neg = any(c.dep_ == "neg" for c in tok.head.children)
            term_text = "not so much" if neg else "so much"
            # Prefer the noun this appositive attaches to (e.g., "Gradescope"), excluding
            # the full "much" subtree so we don't get "Gradescope not much" as the object.
            if tok.head.head is not None and tok.head.head.pos_ in {"NOUN", "PROPN"}:
                obj_text = _object_text_from_noun_head(
                    tok.head.head,
                    evaluative_adjectives=evaluative_adjectives,
                    comparative_words=comparative_words,
                    exclude_tokens=set(tok.head.subtree),
                )
            else:
                obj_text = object_for_modifier(tok.head)
            add_pair(term_text, obj_text)
            continue

        if tok.lemma_.lower() not in evaluative_adverbs:
            continue

        obj_text = object_for_modifier(tok.head)
        add_pair(tok.text, obj_text)

    return pairs
def extract_clauses(text, nlp, *, sent_offset=0):
    """Split text into clauses.

    Returns a list of (sent_num, clause_num, clause_text) where:
    - sent_num is 1-based and sequential within the file (caller supplies sent_offset)
    - clause_num is 1-based within the sentence
    """

    doc = nlp(text)
    clauses = []
    sent_counter = 0
    for sent in doc.sents:
        # Use dependency parse to split on conjunctions and subordinate clauses
        start = 0  # sentence-relative token index

        def _split_idx_for_main_clause_after_preposed_if(verb_token):
            """If verb has a *preposed* 'if' adverbial clause, split before the main clause.

            Example pattern (as parsed by spaCy):
              if ... [advcl] -> head verb, and then main clause starts with a subject + verb.
            We want: "if ..." | "you take ..."
            """

            if verb_token is None or verb_token.pos_ != 'VERB':
                return None

            # Preposed conditional: an advcl child marked by "if" that occurs before the verb.
            preposed_if_advcl = None
            for child in verb_token.children:
                if child.dep_ != 'advcl':
                    continue
                if child.i >= verb_token.i:
                    continue
                if any(c.dep_ == 'mark' and c.lower_ == 'if' for c in child.children):
                    preposed_if_advcl = child
                    break
            if preposed_if_advcl is None:
                return None

            # Prefer splitting at the subject of the main clause (avoids including the advcl in the verb subtree).
            subj = next((c for c in verb_token.children if c.dep_ in {'nsubj', 'nsubjpass'}), None)
            if subj is not None:
                boundary_abs = subj.left_edge.i
                if boundary_abs > preposed_if_advcl.right_edge.i:
                    return boundary_abs - sent.start

            # Fallback: split at the verb itself.
            return (verb_token.i - sent.start)

        clause_counter = 0

        def _emit_clause(end_idx_in_sent):
            nonlocal start, clause_counter
            clause_text = sent[start:end_idx_in_sent].text.strip()
            if clause_text:
                clause_counter += 1
                clauses.append((sent_offset + sent_counter, clause_counter, clause_text))
            start = end_idx_in_sent

        sent_counter += 1

        for token in sent:
            # Split at conjunctions (cc, conj) and markers (mark) for subclauses
            token_idx_in_sent = token.i - sent.start
            adv_class = classify_adv(token) if token.dep_ == 'advmod' else None

            prev_token = sent[token_idx_in_sent - 1] if token_idx_in_sent > 0 else None

            # Avoid dangling coordinators like a standalone "so" when we split at "so" (cc)
            # and the next token is a marker starting a subordinate clause ("so if ...").
            is_mark_following_clause_start_cc = (
                token.dep_ == 'mark'
                and prev_token is not None
                and prev_token.dep_ == 'cc'
                and prev_token.lower_ in {'so'}
                and (token_idx_in_sent - 1) == start
            )

            # Don't split at complementizer "that" for copular constructions like
            # "the problem is that ..."; splitting yields an incomplete clause "the problem is".
            is_copular_complementizer_that = (
                token.dep_ == 'mark'
                and token.lower_ == 'that'
                and token.head.dep_ == 'ccomp'
                and token.head.head.lemma_ == 'be'
            )

            # Treat some consecutive marker patterns as a single unit so we don't
            # create a dangling clause like "because" or "so" on its own.
            #
            # Case A (same head): "so that" (both mark -> same VERB)
            # Case B (nested): "because if ..." / "whereas if ..."
            #   where the second mark (if) attaches to an advcl whose head is the
            #   verb/clause introduced by the first mark (because/whereas).
            is_second_mark_in_pair = False
            if (
                token.dep_ == 'mark'
                and prev_token is not None
                and prev_token.dep_ == 'mark'
                and token.head is not None
                and prev_token.head is not None
            ):
                same_head = prev_token.head.i == token.head.i and token.head.pos_ == 'VERB'

                nested_head = (
                    token.head.head is not None
                    and token.head.dep_ in {'advcl', 'ccomp', 'xcomp'}
                    and token.head.head.i == prev_token.head.i
                    and prev_token.head.pos_ in {'VERB', 'AUX'}
                )

                if same_head or nested_head:
                    is_second_mark_in_pair = True

            # Don't split inside noun-complement constructions like:
            #   "the fact that it's ..." / "the idea that ..."
            # In these, "that" marks an acl modifying a NOUN and should stay attached.
            is_noun_complement_that = (
                token.dep_ == 'mark'
                and token.lower_ == 'that'
                and token.head.dep_ == 'acl'
                and token.head.head is not None
                and token.head.head.pos_ in {'NOUN', 'PROPN'}
            )

            next_tok = None
            if token_idx_in_sent + 1 < len(sent):
                # Find next non-punct token within the sentence
                for j in range(token_idx_in_sent + 1, len(sent)):
                    cand = sent[j]
                    if cand.is_punct or cand.pos_ == 'PUNCT':
                        continue
                    next_tok = cand
                    break

            # New-clause heuristic: split before a main clause that follows a preposed "if" clause.
            # This catches patterns like: "..., like if ..., you take it home ..."
            if token.pos_ == 'VERB' and token_idx_in_sent > start:
                split_idx = _split_idx_for_main_clause_after_preposed_if(token)
                if split_idx is not None and split_idx > start:
                    _emit_clause(split_idx)
                    continue

            # Discourse/conditional "then" often starts a new clause after a comma.
            # Example: "..., then that's all you have."
            if (
                token.dep_ == 'advmod'
                and token.lower_ == 'then'
                and token_idx_in_sent > start
                and token.i > sent.start
                and token.nbor(-1).is_punct
            ):
                _emit_clause(token_idx_in_sent)
                continue

            # Prevent splitting right after a frequency/evidence adverb like “sometimes/obviously”
            # in patterns such as “sometimes if ...” or “obviously if ...”.
            prev_adv_class = None
            if token_idx_in_sent > 0:
                prev_tok = sent[token_idx_in_sent - 1]
                if prev_tok.dep_ == 'advmod':
                    prev_adv_class = classify_adv(prev_tok)
            
            if (
                (
                    (
                        (
                            token.dep_ == 'mark'
                            and not is_second_mark_in_pair
                            and not is_noun_complement_that
                            and not is_mark_following_clause_start_cc
                            and not is_copular_complementizer_that
                        )
                        or (
                            token.dep_ == 'cc'
                            # Only split on coordinators when they join verbal/clausal material.
                            # This avoids fragments like "or something" when "or" coordinates NPs.
                            and (
                                token.head.pos_ in ('VERB', 'AUX')
                                # Also split when the coordinator is followed by a fresh clause start
                                # like "or you VERB ..." even if spaCy attaches the cc under an NP.
                                or (
                                    next_tok is not None
                                    and next_tok.dep_ in ('nsubj', 'nsubjpass')
                                    and next_tok.head.pos_ == 'VERB'
                                )
                            )
                        )
                    )
                    and token_idx_in_sent > start
                    and not (
                        token.dep_ == 'mark'
                        and prev_adv_class in ('FREQUENCY_ADV', 'EVIDENCE_ADV')
                    )
                )
                or (
                    token.dep_ == 'advmod'
                    and adv_class == 'COND_THEN'  # “if ..., then ...” (exclude frequency adverbs)
                    and adv_class not in ('FREQUENCY_ADV', 'EVIDENCE_ADV')
                    and token.head.tag_ == 'VB'
                    and token_idx_in_sent > start
                )
                or (
                    token.dep_ in ('nsubj', 'nsubjpass')
                    and token.pos_ == 'PRON'
                    and token.tag_ in ('WDT', 'WP', 'WP$', 'WRB')
                    and token.head.dep_ in ('advcl','nsubj','nsubjpass','relcl')
                    and token.head.pos_ in ('VERB', 'AUX')
                    and token_idx_in_sent > start
                      # split before new clause starting with a pronoun subject (heuristic for "I think ...", "It was ...")
                )
            ):
                _emit_clause(token_idx_in_sent)
        # Add the last clause
        _emit_clause(len(sent))

    return clauses, (sent_offset + sent_counter)

def get_speaker_lines(doc):
    lines = []
    for para in doc.paragraphs:
        text = re.sub(r"\s+", " ", para.text).strip()
        if text:
            lines.append(text)
    return lines


def get_speaker_utterances(doc):
    """Extract (speaker, utterance_text) tuples from a docx.

    Handles cases where the speaker label is on its own line, e.g.:
      "Speaker 1:"\n"So it sounds like ..."

    Strategy:
    - Detect speaker labels at the start of a paragraph.
    - If a paragraph has only the label, carry the speaker forward.
    - Concatenate subsequent paragraphs for that speaker until the next label.
    """

    # Some transcripts include timestamps like "0:07" either on their own line
    # or preceding the speaker label within the same paragraph.
    timestamp_re = re.compile(r"^\s*\d{1,2}:\d{2}\s*$")
    label_re = re.compile(
        r"^\s*(?:\d{1,2}:\d{2}\s*)?speaker\s*(1|2)\s*:?\s*(.*)$",
        re.IGNORECASE,
    )

    utterances = []
    current_speaker = None
    buffer_parts = []

    def flush():
        nonlocal buffer_parts
        if current_speaker and buffer_parts:
            text = " ".join(buffer_parts).strip()
            if text:
                utterances.append((current_speaker, text))
        buffer_parts = []

    for para in doc.paragraphs:
        raw = para.text
        if not raw or not raw.strip():
            continue

        # Preserve intra-paragraph line breaks so we can detect
        # timestamp lines and speaker labels reliably.
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        if not lines:
            continue

        for line in lines:
            if timestamp_re.match(line):
                continue

            m = label_re.match(line)
            if m:
                flush()
                current_speaker = f"Speaker {m.group(1)}"
                remainder = (m.group(2) or "").strip()
                # Strip leading timestamp that sometimes appears right after the label.
                remainder = re.sub(r"^\d{1,2}:\d{2}\s*", "", remainder)
                if remainder:
                    buffer_parts.append(remainder)
                continue

            # Continuation line: attach to current speaker if known
            if current_speaker:
                buffer_parts.append(re.sub(r"\s+", " ", line).strip())

    flush()
    return utterances

def parse_interview(filename, interviewer, interviewee, nlp):
    doc = docx.Document(filename)

    utterances = get_speaker_utterances(doc)
    data = []
    sent_offset = 0
    for speaker, content in utterances:
        clauses, sent_offset = extract_clauses(content, nlp, sent_offset=sent_offset)
        for sent_num, clause_num, clause_text in clauses:
            eval_pairs = extract_evaluations(clause_text, nlp)
            if eval_pairs:
                eval_terms = "; ".join(t for t, _ in eval_pairs)
                eval_objects = "; ".join(o for _, o in eval_pairs)
            else:
                eval_terms = ""
                eval_objects = ""
            data.append([
                os.path.basename(filename),
                sent_num,
                clause_num,
                speaker,
                clause_text,
                eval_terms,
                eval_objects,
            ])
    return data


def find_word_files(input_dir: Path):
    """Return .docx files in input_dir (non-recursive) sorted alphabetically.

    Skips Word lock/temp files that start with "~$".
    """

    files = []
    for p in input_dir.iterdir():
        if not p.is_file():
            continue
        if p.suffix.lower() != ".docx":
            continue
        if p.name.startswith("~$"):
            continue
        files.append(p)
    return sorted(files, key=lambda x: x.name.lower())

def main():
    input_dir = Path(__file__).resolve().parent
    output_excel = "interview_clauses.xlsx"

    word_files = find_word_files(input_dir)
    if not word_files:
        raise FileNotFoundError(f"No .docx files found in {input_dir}")

    nlp = load_spacy_model()

    # Write to Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Filename", "Sent", "Clause", "Speaker", "Text", "evaluative term", "object"])

    total_rows = 0
    for word_path in word_files:
        rows = parse_interview(str(word_path), None, None, nlp)
        for row in rows:
            ws.append(row)
        total_rows += len(rows)
        print(f"Processed {word_path.name}: {len(rows)} rows")

    wb.save(output_excel)
    print(f"Saved {total_rows} rows from {len(word_files)} files to {output_excel}")

if __name__ == '__main__':
    main()
